# -*- coding: utf-8 -*-
"""One-off: bring Bullyproof-SOW-Completion-Register-FINAL.docx in line with
the house style - full-bleed cover page and every numbered top-level section
starting on a new page. The body text is left untouched."""
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Emu

conv = __import__("convert-handover-docs")

DOC = Path(__file__).resolve().parents[1] / "docs" / "handover" / "Bullyproof-SOW-Completion-Register-FINAL.docx"
COVER = conv.COVERS_DIR / "sow-completion-register.png"

doc = Document(str(DOC))

# 1. Numbered top-level sections open a fresh page.
broken = []
for p in doc.paragraphs:
    if p.style is not None and p.style.name == "Heading 1" and re.match(r"^\d+[a-z]?[\.\)]\s*", p.text.strip()):
        p.paragraph_format.page_break_before = True
        broken.append(p.text.strip()[:60])

# 2. Full-bleed cover as page one; existing letterhead block moves to page 2.
first = doc.paragraphs[0]
cover_par = first.insert_paragraph_before()
run = cover_par.add_run()
run.add_picture(str(COVER), width=Emu(conv.A4_WIDTH_EMU), height=Emu(conv.A4_HEIGHT_EMU))
conv.anchor_picture_to_page(run)
run.add_break(WD_BREAK.PAGE)

# 3. Keep the cover page clean: no header/footer on page one.
section = doc.sections[0]
section.different_first_page_header_footer = True

doc.save(str(DOC))
print("page-break sections:")
for b in broken:
    print("  -", b)
print("cover inserted, saved", DOC.name)
