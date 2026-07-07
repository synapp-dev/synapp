"""
Converts the handover markdown documents (deliverable D6 and the acceptance
paperwork) to polished Word .docx for the client package.

Usage (from apps/bullyproof):
    python scripts/convert-handover-docs.py docs/handover/admin-user-guide.md ...

Writes a .docx next to each input with the same basename. Handles the
markdown subset the handover docs use: headings, paragraphs, bullet and
numbered lists (one nesting level), GFM pipe tables, fenced code blocks,
horizontal rules, bold / italic / inline code / links.

House style: the Intradark wordmark on the cover and in the running header,
ink headings with a teal accent (the client's corporate teal), numbered
top-level sections always start on a new page, teal table headers, and a
branded footer with page numbers.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

INK = RGBColor(0x1F, 0x29, 0x37)
TEAL = RGBColor(0x00, 0x84, 0x90)
GREY = RGBColor(0x6B, 0x72, 0x80)
CODE_INK = RGBColor(0x3E, 0x45, 0x4E)
TEAL_HEX = "008490"

FOOTER_TEXT = "Intradark Pty Ltd  ·  ABN 38 696 182 457  ·  Commercial in confidence"

WORDMARK = Path(__file__).resolve().parent / "assets" / "intradark-wordmark.png"
# Sizes lifted from the 29 June proposal set so the whole package matches.
WORDMARK_COVER_WIDTH = Emu(1952625)  # ~5.4 cm on the first page
WORDMARK_HEADER_WIDTH = Emu(904875)  # ~2.5 cm in the running header

INLINE_TOKEN = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
NUMBERED_SECTION = re.compile(r"^\d+[a-z]?[\.\)]\s")


def add_inline(paragraph, text):
    """Render markdown inline formatting into runs on the paragraph."""
    for part in INLINE_TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = CODE_INK
        else:
            match = LINK.fullmatch(part)
            if match:
                label, url = match.group(1), match.group(2)
                run = paragraph.add_run(label)
                run.bold = True
                if url != label:
                    paragraph.add_run(f" ({url})")
            else:
                paragraph.add_run(part)


def set_paragraph_border(paragraph, edge, color_hex, size=6, space=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    element = OxmlElement(f"w:{edge}")
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), str(size))
    element.set(qn("w:space"), str(space))
    element.set(qn("w:color"), color_hex)
    p_bdr.append(element)


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_number_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    return run


def fill_footer(footer, section):
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = ""
    tab_stops = footer_paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(
        section.page_width - section.left_margin - section.right_margin,
        WD_TAB_ALIGNMENT.RIGHT,
    )
    run = footer_paragraph.add_run(FOOTER_TEXT)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    footer_paragraph.add_run("\t").font.size = Pt(8)
    page_label = footer_paragraph.add_run("Page ")
    page_label.font.size = Pt(8)
    page_label.font.color.rgb = GREY
    add_page_number_field(footer_paragraph, "PAGE")
    of_label = footer_paragraph.add_run(" of ")
    of_label.font.size = Pt(8)
    of_label.font.color.rgb = GREY
    add_page_number_field(footer_paragraph, "NUMPAGES")
    set_paragraph_border(footer_paragraph, "top", TEAL_HEX, size=4, space=6)


def build_branding(doc):
    """Footer on every page; wordmark in the header of every page but the
    first (the first page carries the large cover wordmark in the body)."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    fill_footer(section.footer, section)
    fill_footer(section.first_page_footer, section)

    if WORDMARK.exists():
        header_paragraph = section.header.paragraphs[0]
        header_paragraph.text = ""
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_paragraph.add_run().add_picture(
            str(WORDMARK), width=WORDMARK_HEADER_WIDTH
        )


def add_cover_wordmark(doc):
    if not WORDMARK.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(18)
    paragraph.add_run().add_picture(str(WORDMARK), width=WORDMARK_COVER_WIDTH)


def style_heading(paragraph, level, text_is_numbered):
    for run in paragraph.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(25)
            run.font.color.rgb = INK
            run.bold = True
        elif level == 2:
            run.font.size = Pt(15.5)
            run.font.color.rgb = TEAL
            run.bold = True
        elif level == 3:
            run.font.size = Pt(12.5)
            run.font.color.rgb = INK
            run.bold = True
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = GREY
            run.bold = True
    fmt = paragraph.paragraph_format
    if level == 1:
        fmt.space_after = Pt(4)
        set_paragraph_border(paragraph, "bottom", TEAL_HEX, size=10, space=8)
    elif level == 2:
        fmt.space_before = Pt(16)
        fmt.space_after = Pt(6)
    else:
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(4)
    # Numbered top-level sections always open a fresh page.
    if level in (2, 3) and text_is_numbered:
        fmt.page_break_before = True


def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def convert(md_path: Path) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = INK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.12

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    build_branding(doc)
    add_cover_wordmark(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing fence
            for code_line in code_lines or [""]:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Pt(18)
                paragraph.paragraph_format.space_after = Pt(0)
                run = paragraph.add_run(code_line if code_line else " ")
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = CODE_INK
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            heading_text = heading.group(2).strip()
            paragraph = doc.add_heading("", level=min(level, 4))
            add_inline(paragraph, heading_text)
            style_heading(
                paragraph, level, bool(NUMBERED_SECTION.match(heading_text))
            )
            i += 1
            continue

        if re.fullmatch(r"[-*_]{3,}", stripped):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(8)
            set_paragraph_border(paragraph, "bottom", TEAL_HEX, size=4, space=1)
            i += 1
            continue

        if is_table_row(stripped) and i + 1 < len(lines) and re.match(
            r"^\|[\s:|-]+\|$", lines[i + 1].strip()
        ):
            header = split_row(stripped)
            i += 2
            body = []
            while i < len(lines) and is_table_row(lines[i]):
                body.append(split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            for col, text in enumerate(header):
                cell = table.rows[0].cells[col]
                cell.text = ""
                shade_cell(cell, TEAL_HEX)
                paragraph = cell.paragraphs[0]
                add_inline(paragraph, text)
                if not paragraph.runs:
                    paragraph.add_run(" ")
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for row_index, row_values in enumerate(body):
                row = table.add_row()
                for col in range(len(header)):
                    value = row_values[col] if col < len(row_values) else ""
                    cell = row.cells[col]
                    cell.text = ""
                    if row_index % 2 == 1:
                        shade_cell(cell, "F2F7F7")
                    paragraph = cell.paragraphs[0]
                    add_inline(paragraph, value)
                    for run in paragraph.runs:
                        if run.font.size is None:
                            run.font.size = Pt(9.5)
            doc.add_paragraph()
            continue

        bullet = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if bullet:
            level = 2 if len(bullet.group(1)) >= 2 else 1
            style_name = "List Bullet" if level == 1 else "List Bullet 2"
            paragraph = doc.add_paragraph(style=style_name)
            add_inline(paragraph, bullet.group(2).strip())
            i += 1
            continue

        numbered = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if numbered:
            level = 2 if len(numbered.group(1)) >= 2 else 1
            style_name = "List Number" if level == 1 else "List Number 2"
            paragraph = doc.add_paragraph(style=style_name)
            add_inline(paragraph, numbered.group(2).strip())
            i += 1
            continue

        if stripped.startswith(">"):
            paragraph = doc.add_paragraph()
            add_inline(paragraph, stripped.lstrip("> "))
            paragraph.paragraph_format.left_indent = Pt(14)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(10)
            set_paragraph_border(paragraph, "left", TEAL_HEX, size=12, space=8)
            for run in paragraph.runs:
                run.font.color.rgb = GREY
                run.font.size = Pt(10.5)
                run.italic = True
            i += 1
            continue

        # Plain paragraph: merge soft-wrapped continuation lines.
        text_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt.startswith(("#", "```", "|", "- ", "* ", ">"))
                or re.match(r"^\d+\.\s", nxt)
                or re.fullmatch(r"[-*_]{3,}", nxt)
            ):
                break
            text_lines.append(nxt)
            i += 1
        paragraph = doc.add_paragraph()
        add_inline(paragraph, " ".join(text_lines))

    out_path = md_path.with_suffix(".docx")
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    for arg in sys.argv[1:]:
        path = Path(arg)
        out = convert(path)
        print(f"{path} -> {out}")
