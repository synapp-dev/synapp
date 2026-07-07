"""
Builds the FINAL issue of the SOW Completion & Deliverables Register from the
29 June original, flipping the remaining "Finishing" rows to Delivered with
updated evidence, rewriting Sections 2 and 3 in the past tense, and updating
the date and closing.

Usage (from apps/bullyproof):
    python scripts/build-final-completion-register.py

Reads  docs/change-request/Bullyproof-SOW-Completion-Register.docx
Writes docs/handover/Bullyproof-SOW-Completion-Register-FINAL.docx
"""

from copy import deepcopy
from pathlib import Path

import docx

SRC = Path("docs/change-request/Bullyproof-SOW-Completion-Register.docx")
OUT = Path("docs/handover/Bullyproof-SOW-Completion-Register-FINAL.docx")

FINAL_DATE = "7 July 2026"


def set_text(paragraph, text):
    """Replace a paragraph's text, keeping the first run's formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell(cell, text):
    set_text(cell.paragraphs[0], text)
    for extra in cell.paragraphs[1:]:
        set_text(extra, "")


# (first-cell match, new status or None, new evidence or None)
ROW_EDITS = [
    (
        "~60-day OTP expiry (ST4S)",
        "Delivered",
        "supabase/config.toml (otp_expiry); session policy documented in "
        "docs/handover/auth-session-expiry.md",
    ),
    (
        "Certificate on completion, in profile",
        None,
        "TopicCertificate; markCertificateIssued; downloadable PDF on the "
        "official certificate artwork via /api/certification/courses/[id]/certificate; "
        "Certificates section on /profile",
    ),
    (
        "Government view-only dashboard",
        "Delivered",
        "GovernmentDashboard: five platform-wide aggregates, view-only, "
        "CSV/PDF export; /api/government/overview; verified live 7 Jul 2026",
    ),
    (
        "CSV / PDF export",
        "Delivered",
        "Export menu (CSV + PDF) on all four admin report tabs plus "
        "role-scoped export packs for school admins and teachers on the "
        "school Reports page (lib/report-export.ts; server/reports)",
    ),
    (
        "Calculations incl. zero-input weighting",
        None,
        "lib/culture-rating-math.ts; zero-input re-weighting unit-tested and "
        "verified live 7 Jul 2026 (official Woodford constants applied on "
        "receipt of your mapping)",
    ),
    (
        "D6 — Documentation set (admin / user / tech)",
        "Delivered",
        "docs/handover: Administrator User Guide and System Administrator "
        "Guide (Word copies supplied with this register)",
    ),
    (
        "D7 — Bill of Materials (OSS licences)",
        "Delivered",
        "docs/handover/bill-of-materials.md and .csv: 640 packages, "
        "generated from the production lockfile at final delivery",
    ),
]

SECTION2_HEADING = "2.  Previously remaining SOW deliverables, now complete"
SECTION2_INTRO = (
    "The items marked Finishing in the 29 June issue of this register are now "
    "delivered, at no additional charge, as promised:"
)
SECTION2_BULLETS = {
    "Government view-only reporting dashboard": (
        "Government view-only reporting dashboard: delivered. A live "
        "aggregate dashboard (schools, licences, staff, lessons delivered, "
        "AP certified) with CSV/PDF export, view-only by design."
    ),
    "CSV / PDF export across the reporting roles": (
        "CSV / PDF export across the reporting roles: delivered. Export on "
        "all four admin report tabs, plus role-scoped export packs: the "
        "schools register for Bullyproof admin, class / staff / lesson "
        "history / culture packs per school, and a personal pack for "
        "teachers on the school Reports page."
    ),
    "Downloadable AP certificate artifact": (
        "Downloadable AP certificate artifact: delivered. The certificate "
        "renders on the official Certificate of Completion artwork and "
        "downloads from the course page and the user profile."
    ),
    "Administrator / user / technical documentation": (
        "Administrator / user / technical documentation and the Bill of "
        "Materials: delivered with this register (Administrator User Guide, "
        "System Administrator Guide, bill-of-materials.md/.csv)."
    ),
    "Final culture-rating weighting": (
        "Final culture-rating weighting: the calculation engine is delivered, "
        "unit-tested and verified live, including weight re-distribution when "
        "a metric has no measurable change. The official Woodford constants "
        "are applied the day you supply the mapping; this is the one open "
        "input and it does not affect acceptance of the delivered system."
    ),
    "OTP expiry set to ~60 days (ST4S)": (
        "OTP expiry (ST4S): configured, with the full session-expiry position "
        "documented in docs/handover/auth-session-expiry.md."
    ),
}

SECTION3_EDITS = {
    "School Address & Email Domain not saving": (
        "School Address & Email Domain not saving on edit (39): reproduced, "
        "root-caused and fixed. The save always persisted; the admin school "
        "drawer then failed to display the saved values, which reads as a "
        "failed save. The display fault is fixed and the exact flow verified: "
        "edit, save, reload, values shown."
    ),
    "Wrong recommended lesson for mixed-level classes": (
        "Wrong recommended lesson for mixed-level classes, and the list "
        "capped at three (53): fixed. Composite classes resolve to the "
        "lowest matching stage with a per-class guidance panel, and the "
        "topic list is uncapped."
    ),
    "Teach Lessons — blank screen": (
        "Teach Lessons blank screen and the Back to lesson loop (51): fixed. "
        "A clear fallback panel renders when no recommendation applies, and "
        "the return path routes by lesson status."
    ),
    "Apply Template gives no confirmation": (
        "Apply Template gives no confirmation it applied (40): fixed, and "
        "beyond the promised toast the template cards now show a persistent "
        "Active on N schools status."
    ),
}

SECTION3_EXTRA = (
    "Also found and fixed during completion UAT, at no cost: user-table "
    "column sorting across pages (28), a fault that blocked creating "
    "platform-scoped users, and a welcome-tutorial redirect loop when "
    "impersonating a school user. Beyond the defects, the minor text and "
    "usability asks from your 22 June document have been applied at no "
    "charge as promised: only the highest access level shows per user, the "
    "sidebar screen tips use your wording, the Teach Lessons page and "
    "lesson-wizard copy follow your notes (including removal of the step "
    "numbers, search bar, Show completed and Help), the feedback form "
    "carries the mandatory-rating note, and the school portal filters read "
    "All Access Levels."
)

CLOSING = (
    "In short: every deliverable under the Agreement and the Variation is "
    "Delivered. This register, the documentation set, the Bill of Materials "
    "and the UAT checklist accompany the Final Acceptance letter. Acceptance "
    "is requested under Variation clause 7, and the clause 4(b) completion "
    "payment falls due on acceptance and handover."
)


def main():
    doc = docx.Document(SRC)

    # Header table: date and subject line.
    header = doc.tables[0]
    for row in header.rows:
        label = row.cells[0].text.strip()
        if label == "Date:":
            set_cell(row.cells[1], FINAL_DATE)
        elif label == "Re:":
            set_cell(
                row.cells[1],
                "FINAL ISSUE: completion of all deliverables under the Formal "
                "Variation Agreement and Statement of Work; acceptance "
                "requested under clause 7",
            )

    # Deliverables table: flip statuses and refresh evidence.
    deliverables = doc.tables[1]
    edited = set()
    for row in deliverables.rows:
        first = row.cells[0].text.strip()
        for match, status, evidence in ROW_EDITS:
            if first.startswith(match[:20]):
                if status:
                    set_cell(row.cells[2], status)
                if evidence:
                    set_cell(row.cells[3], evidence)
                edited.add(match)
                break
    missing = [m for m, _, _ in ROW_EDITS if m not in edited]
    if missing:
        raise SystemExit(f"Rows not found in deliverables table: {missing}")

    section3_anchor = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text.startswith("Status is one of Delivered or Finishing"):
            set_text(
                paragraph,
                f"As at {FINAL_DATE} every deliverable is Delivered. Items "
                "previously marked Finishing in the 29 June issue are now "
                "complete; Section 2 records each one. Evidence cites the "
                "implementing code.",
            )
        elif text.startswith("2.  Remaining SOW deliverables"):
            set_text(paragraph, SECTION2_HEADING)
        elif text.startswith("For completeness, the items above marked"):
            set_text(paragraph, SECTION2_INTRO)
        elif text.startswith("Naming these plainly is deliberate"):
            set_text(
                paragraph,
                "Naming these plainly was deliberate, and the promise has "
                "been kept: the Agreement is honoured in full, and the line "
                "below stands on that honesty.",
            )
        elif text.startswith("I fix genuine defects for free"):
            set_text(
                paragraph,
                "I fix genuine defects for free: a defect being a fault in "
                "functionality the Agreement required and I delivered (SDA "
                "Warranty cl 8.1(b), 90-day cover; Variation cl 7.3). The "
                "following were genuine and have been fixed at no cost:",
            )
        elif text.startswith("In short: the original Agreement"):
            set_text(paragraph, CLOSING)
        else:
            for prefix, replacement in SECTION2_BULLETS.items():
                if text.startswith(prefix):
                    set_text(paragraph, replacement)
                    break
            else:
                for prefix, replacement in SECTION3_EDITS.items():
                    if text.startswith(prefix[:30]):
                        set_text(paragraph, replacement)
                        if prefix.startswith("Apply Template"):
                            section3_anchor = paragraph
                        break

    if section3_anchor is None:
        raise SystemExit("Apply Template bullet not found for Section 3 anchor")

    extra = deepcopy(section3_anchor._p)
    section3_anchor._p.addnext(extra)
    extra_paragraph = docx.text.paragraph.Paragraph(extra, section3_anchor._parent)
    set_text(extra_paragraph, SECTION3_EXTRA)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
