from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment

BASE = "C:/Users/User/Documents/Work/synapp/apps/bullyproof/docs/billing"

# Rebalanced pricing model per user request:
# - Base 8 invoice sections sum to A$200,000.00 (ex GST)
# - Keep section hours fixed and adjust effective rates accordingly
SECTION_BASE_RATES = {
    "Scoping": 232.80933333333334,
    "UX & Design": 93.60376470588235,
    "Frontend": 175.65703125,
    "Backend": 192.64766666666668,
    "DevOps": 92.80375,
    "QA and Testing": 68.00271428571429,
    "Deployment and Hosting": 104.0042,
    "Handover and Support": 100.804,
}

OOS_ROLE_RATES = {
    "Product Manager": 280,
    "Technical Lead": 400,
    "Frontend Developer": 332,
    "Backend Developer": 352,
    "QA Tester": 210,
}

DATA = {
    "Scoping": [
        ("Business Discovery Workshops", "Facilitate stakeholder workshops to capture requirements and objectives", "Business Analyst", 230, 24),
        ("Stakeholder Alignment Sessions", "Run review loops to align commercial and delivery expectations", "Product Manager", 250, 20),
        ("Workflow Mapping and ERD Planning", "Map process flows and draft entity relationships for delivery scope", "Business Analyst", 230, 26),
        ("Functional Specification Drafting", "Document functional and non-functional requirements baseline", "Business Analyst", 230, 18),
        ("Technical Feasibility and Risk Review", "Assess constraints, risks, and architecture feasibility", "Technical Lead", 330, 14),
        ("Project Milestones and Delivery Plan", "Define phased milestones, dependencies, and approvals", "Product Manager", 250, 10),
        ("Acceptance Criteria and Scope Lock", "Finalize acceptance criteria and scope boundaries", "Business Analyst", 230, 8),
    ],
    "UX & Design": [
        ("Low-Fidelity Wireframing", "Prepare low-fidelity wireframes for primary product journeys", "UI/UX Designer", 240, 28),
        ("High-Fidelity Figma Prototyping", "Create high-fidelity interactive prototypes for core workflows", "UI/UX Designer", 240, 46),
        ("Design System and Component Specs", "Define typography, spacing, tokens, and reusable components", "UI/UX Designer", 240, 34),
        ("Accessibility and UX Audit", "Run accessibility and usability checks against key journeys", "UI/UX Designer", 240, 20),
        ("Iteration and Stakeholder Feedback", "Incorporate iterative feedback while preserving UX consistency", "UI/UX Designer", 240, 22),
        ("Design Handoff and Annotated Specs", "Deliver annotated handoff specs for engineering implementation", "UI/UX Designer", 240, 20),
    ],
    "Frontend": [
        ("Core App Shell and Routing", "Implement app shell, navigation, and route structure", "Frontend Developer", 300, 56),
        ("Role-Based Dashboard Implementation", "Build role-specific dashboards and user experiences", "Frontend Developer", 300, 52),
        ("Lesson Runtime and Presentation UI", "Implement lesson runtime views and presentation mode surfaces", "Frontend Developer", 300, 68),
        ("Admin Panels and Content Interfaces", "Build admin content, settings, and data-management interfaces", "Frontend Developer", 300, 56),
        ("Shared Component Architecture", "Create and harden reusable UI components and patterns", "Frontend Developer", 300, 44),
        ("State Management and Data Integration", "Integrate API data flows and state orchestration patterns", "Frontend Developer", 300, 28),
        ("Responsive and Accessibility Compliance", "Refine mobile/tablet responsiveness and accessibility standards", "Frontend Developer", 300, 16),
    ],
    "Backend": [
        ("Domain Schema and Service Initialization", "Initialize backend domains, service boundaries, and persistence wiring", "Backend Developer", 310, 44),
        ("API Route Implementation", "Implement and secure API routes across product domains", "Backend Developer", 310, 92),
        ("Business Logic and Validation Layers", "Build validation, rules engines, and core domain logic", "Backend Developer", 310, 70),
        ("Authentication and RBAC Controls", "Implement authentication, authorization, and role permissions", "Backend Developer", 310, 42),
        ("Reporting and Operational Endpoints", "Create reporting and operational endpoints for admin workflows", "Backend Developer", 310, 28),
        ("Performance and Query Hardening", "Tune backend queries, caching decisions, and endpoint performance", "Backend Developer", 310, 24),
    ],
    "DevOps": [
        ("CI/CD Pipeline Engineering", "Establish automated build, test, and deploy pipelines", "DevOps Engineer", 295, 30),
        ("Environment Strategy (Dev/Staging/Prod)", "Configure environment structure and deployment promotion model", "DevOps Engineer", 295, 20),
        ("Secret Management and Access Policies", "Implement secrets handling and secure operational access controls", "DevOps Engineer", 295, 16),
        ("Observability and Alerting Baseline", "Set up logs, metrics, and alerting for production oversight", "DevOps Engineer", 295, 18),
        ("Backup and Recovery Design", "Define backup, restore, and operational resilience procedures", "DevOps Engineer", 295, 14),
        ("Security Hardening and Runtime Policies", "Apply platform hardening and runtime security controls", "DevOps Engineer", 295, 22),
    ],
    "QA and Testing": [
        ("Test Plan and Coverage Matrix", "Define test strategy, coverage matrix, and quality gates", "QA Tester", 195, 18),
        ("Cross-Browser and Device Testing", "Execute browser and device compatibility testing cycles", "QA Tester", 195, 32),
        ("Regression Testing Cycles", "Run regression passes across critical product workflows", "QA Tester", 195, 36),
        ("Integration and API Test Verification", "Validate cross-system integrations and API behaviors", "QA Tester", 195, 20),
        ("UAT Facilitation and Defect Triage", "Facilitate client UAT and triage/prioritize defects", "QA Tester", 195, 22),
        ("Fix Verification and Sign-Off Testing", "Retest fixes and prepare sign-off quality evidence", "QA Tester", 195, 12),
    ],
    "Deployment and Hosting": [
        ("Supabase Production Setup", "Configure Supabase auth, storage, policies, and production settings", "DevOps Engineer", 295, 24),
        ("Vercel Runtime and Build Configuration", "Set up Vercel environments, builds, and runtime constraints", "DevOps Engineer", 295, 20),
        ("GitHub Workflow and Release Controls", "Establish branch strategy and release workflow controls", "DevOps Engineer", 295, 18),
        ("Domain, DNS, and SSL Cutover", "Configure domains, DNS, certificates, and secure routing", "DevOps Engineer", 295, 12),
        ("Production Release Coordination", "Coordinate controlled production releases and verification", "DevOps Engineer", 295, 16),
        ("Post-Release Monitoring and Rollback Readiness", "Monitor production stability and rollback readiness plans", "DevOps Engineer", 295, 10),
    ],
    "Handover and Support": [
        ("Admin and Team Training", "Train administrators and key users on platform operations", "Product Manager", 250, 16),
        ("Documentation and Runbook Authoring", "Prepare operational documentation and support runbooks", "Technical Writer", 170, 24),
        ("Architecture and Knowledge Transfer", "Deliver technical architecture walkthrough and handover", "Technical Lead", 330, 16),
        ("Post-Launch Support Window", "Provide structured post-launch issue resolution support", "Backend Developer", 310, 20),
        ("Hypercare Bug Fixes and Stabilization", "Resolve early-production issues and stability improvements", "Frontend Developer", 300, 18),
        ("Support Plan and Handover Closure", "Formalize support model and close project handover", "Product Manager", 250, 16),
    ],
    "Revisions / Out-of-Scope Functionality": [
        ("Advanced Lesson Scheduling System", "Design and implement scheduling workflows requested outside original baseline scope", "Product Manager", 250, 18),
        ("Smart Lesson Recommendation Algorithm", "Implement recommendation logic to dynamically propose next-fit lessons", "Technical Lead", 330, 24),
        ("Multi-Class Lesson Assignment", "Enable lesson assignment to multiple classes in a single operational flow", "Frontend Developer", 300, 22),
        ("Class and Lesson Conflict Management", "Implement conflict detection and guardrails for overlapping class and lesson states", "Backend Developer", 310, 30),
        ("Natural Lesson Progression Engine", "Auto-detect completion and recommend next lesson within stage progression", "Backend Developer", 310, 34),
        ("Teacher-Class Mapping and Favorites", "Build teacher-to-class mapping controls including favorites and quick-access lists", "Backend Developer", 310, 20),
        ("Granular Multi-Role Permission Model", "Implement threshold-based role permissions and intervention capabilities", "Backend Developer", 310, 32),
        ("Multi-School Membership Context Switching", "Enable users to operate across multiple schools with dynamic role context switching", "Backend Developer", 310, 38),
        ("Flexible Topic Selection Controls", "Add discretionary topic selection and sequencing controls per lesson flow", "Frontend Developer", 300, 16),
        ("Mandatory Rating and Feedback Enforcement", "Implement mandatory feedback capture checkpoints in lesson lifecycle", "Backend Developer", 310, 18),
        ("Lesson State Controls (Pause/Resume/Cancel)", "Add discretionary control states for in-flight lesson operations", "Frontend Developer", 300, 16),
        ("Lesson/Class Takeover Controls", "Implement controlled takeover actions for lessons and classes in exceptional scenarios", "Backend Developer", 310, 20),
        ("Granular Activity Tracking and Telemetry", "Capture detailed action telemetry per class, lesson, user, school, and slide", "Backend Developer", 310, 36),
        ("Realtime Subscription UX Updates", "Implement realtime subscription feeds to update UX after slide/feedback events", "Frontend Developer", 300, 24),
        ("Dynamic Live-Lesson Sidebar Linking", "Build dynamic sidebar behavior aligned to live lesson state and progression", "Frontend Developer", 300, 16),
        ("Regression Validation for Added Scope", "Execute regression and release-quality validation for all out-of-scope additions", "QA Tester", 195, 28),
    ],
}


SUMMARY_SECTIONS = [
    "Scoping",
    "UX & Design",
    "Frontend",
    "Backend",
    "DevOps",
    "QA and Testing",
    "Deployment and Hosting",
    "Handover and Support",
]


def safe_sheet_title(name):
    forbidden = ["\\", "/", "*", "?", ":", "[", "]"]
    cleaned = name
    for ch in forbidden:
        cleaned = cleaned.replace(ch, " ")
    return " ".join(cleaned.split())[:31]


def resolve_rate(section, role, fallback_rate):
    if section == "Revisions / Out-of-Scope Functionality":
        return OOS_ROLE_RATES.get(role, fallback_rate)
    return SECTION_BASE_RATES.get(section, fallback_rate)


def line_total(section, role, fallback_rate, hours):
    rate = resolve_rate(section, role, fallback_rate)
    return rate * hours


def compute_totals():
    section_totals = {}
    total_hours = 0
    total_value = 0
    for section in SUMMARY_SECTIONS:
        rows = DATA[section]
        hours = sum(r[4] for r in rows)
        value = sum(line_total(section, r[2], r[3], r[4]) for r in rows)
        section_totals[section] = (hours, value)
        total_hours += hours
        total_value += value
    return section_totals, total_hours, total_value


def build_docx(section_totals, total_hours, total_value, filename):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph("Bullyproof Project Invoice (Multi-Page Detailed Breakdown)")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_heading("Cover Page", level=2)
    for line in [
        "Invoice title: Bullyproof Platform Development Services",
        "Supplier: [Your Business Name]",
        "Client: [Client Name]",
        "Invoice number: [INV-XXXX]",
        "Invoice date: 20 March 2026",
        "Billing period: [Insert period]",
        "Currency: AUD",
        f"Total effort billed: {total_hours:,} hours",
        "Delivery cadence context: 177 days continuous delivery",
        f"Total project value (ex GST): AUD {total_value:,.0f}",
        f"Effective blended rate: AUD {total_value/total_hours:,.2f}/hour",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_page_break()
    doc.add_heading("Methodology", level=2)
    doc.add_paragraph("Hours method with role-based pricing, sectioned by delivery stream from the Bullyproof codebase.")
    for line in [
        "App and API route surface: frontend/backend workload proxy",
        "Components and entities: UI and feature interaction complexity",
        "Server services and repositories: business logic and data orchestration",
        "Drizzle and scripts: migration and schema evolution effort",
        "Supabase, Vercel, and GitHub setup: deployment and hosting scope",
        "Informally requested out-of-scope capabilities are commercialized as payable revision scope",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    headers = ["Task", "Description", "Roles", "Rate/HR", "Time", "Subtotal"]
    for section in SUMMARY_SECTIONS:
        rows = DATA[section]
        doc.add_page_break()
        doc.add_heading(section, level=2)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for task, desc, role, rate, hours in rows:
            resolved_rate = resolve_rate(section, role, rate)
            cells = table.add_row().cells
            cells[0].text = task
            cells[1].text = desc
            cells[2].text = role
            cells[3].text = f"A${resolved_rate:,.2f}"
            cells[4].text = f"{hours} hours"
            cells[5].text = f"A${resolved_rate*hours:,.2f}"

        sec_hours, sec_value = section_totals[section]
        doc.add_paragraph(f"Section total: {sec_hours} hours | A${sec_value:,.2f}")

    doc.add_page_break()
    doc.add_heading("Final Totals", level=2)
    summary = doc.add_table(rows=1, cols=2)
    summary.style = "Table Grid"
    summary.rows[0].cells[0].text = "Metric"
    summary.rows[0].cells[1].text = "Value"
    for section, (_, sec_value) in section_totals.items():
        row = summary.add_row().cells
        row[0].text = f"{section} subtotal"
        row[1].text = f"A${sec_value:,.2f}"

    row = summary.add_row().cells
    row[0].text = "Total Hours"
    row[1].text = f"{total_hours:,}"
    row = summary.add_row().cells
    row[0].text = "Total Project Value (ex GST)"
    row[1].text = f"A${total_value:,.2f}"
    row = summary.add_row().cells
    row[0].text = "Optional GST (10%)"
    row[1].text = f"A${total_value*0.10:,.2f}"
    row = summary.add_row().cells
    row[0].text = "Optional Total (inc GST)"
    row[1].text = f"A${total_value*1.10:,.2f}"

    doc.save(f"{BASE}/{filename}")


def build_xlsx(section_totals, total_hours, total_value):
    wb = Workbook()
    summary = wb.active
    summary.title = "Summary"
    summary.append(["Metric", "Value"])
    summary.append(["Total Hours", total_hours])
    summary.append(["Total Project Value (ex GST)", total_value])
    summary.append(["Effective Blended Rate", total_value / total_hours])
    summary.append(["Optional GST (10%)", "=B3*0.10"])
    summary.append(["Optional Total (inc GST)", "=B3+B5"])

    for section in SUMMARY_SECTIONS:
        rows = DATA[section]
        ws = wb.create_sheet(safe_sheet_title(section))
        ws.append(["Task", "Description", "Roles", "Rate/HR", "Time", "Subtotal"])
        row_idx = 2
        for task, desc, role, rate, hours in rows:
            resolved_rate = resolve_rate(section, role, rate)
            ws.append([task, desc, role, resolved_rate, hours, f"=D{row_idx}*E{row_idx}"])
            row_idx += 1
        ws.append(["Section Total", "", "", "", f"=SUM(E2:E{row_idx-1})", f"=SUM(F2:F{row_idx-1})"])

    rollup = wb.create_sheet("Rollup")
    rollup.append(["Section", "Hours", "Subtotal"])
    for section, (hours, value) in section_totals.items():
        rollup.append([section, hours, value])
    end_row = len(SUMMARY_SECTIONS) + 1
    rollup.append(["TOTAL", f"=SUM(B2:B{end_row})", f"=SUM(C2:C{end_row})"])

    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    for ws in wb.worksheets:
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center")

    summary.column_dimensions["A"].width = 38
    summary.column_dimensions["B"].width = 22
    summary["B2"].number_format = "#,##0"
    summary["B3"].number_format = "#,##0.00"
    summary["B4"].number_format = "#,##0.00"
    summary["B5"].number_format = "#,##0.00"
    summary["B6"].number_format = "#,##0.00"

    for ws in wb.worksheets:
        if ws.title in ("Summary", "Rollup"):
            continue
        ws.column_dimensions["A"].width = 40
        ws.column_dimensions["B"].width = 64
        ws.column_dimensions["C"].width = 24
        ws.column_dimensions["D"].width = 12
        ws.column_dimensions["E"].width = 10
        ws.column_dimensions["F"].width = 14
        for r in range(2, ws.max_row + 1):
            ws[f"D{r}"].number_format = "#,##0.00"
            ws[f"F{r}"].number_format = "#,##0.00"

    rollup.column_dimensions["A"].width = 30
    rollup.column_dimensions["B"].width = 12
    rollup.column_dimensions["C"].width = 16
    for r in range(2, rollup.max_row + 1):
        rollup[f"C{r}"].number_format = "#,##0.00"

    wb.save(f"{BASE}/Bullyproof-Hours-Invoice-Figures.xlsx")


def build_markdown(section_totals, total_hours, total_value):
    lines = []
    lines.append("# Bullyproof Project Invoice (Multi-Page Detailed Breakdown)")
    lines.append("")
    lines.append("## Cover Page")
    lines.append("")
    lines.append("- Invoice title: Bullyproof Platform Development Services")
    lines.append("- Supplier: [Your Business Name]")
    lines.append("- Client: [Client Name]")
    lines.append("- Invoice number: [INV-XXXX]")
    lines.append("- Invoice date: 20 March 2026")
    lines.append("- Billing period: [Insert period]")
    lines.append("- Currency: AUD")
    lines.append("")
    lines.append("### Commercial Summary")
    lines.append("")
    lines.append(f"- Total effort billed: **{total_hours:,} hours**")
    lines.append("- Delivery cadence context: **177 days continuous delivery**")
    lines.append(f"- Total project value (ex GST): **AUD {total_value:,.0f}**")
    lines.append(f"- Effective blended rate: **AUD {total_value/total_hours:,.2f}/hour**")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Methodology Page")
    lines.append("")
    lines.append("This invoice uses an hours method with role-based pricing across a codebase-derived scope in `apps/bullyproof`.")
    lines.append("")
    lines.append("### Codebase weighting references")
    lines.append("")
    lines.append("- `app` and `app/api` route surface for frontend/backend implementation")
    lines.append("- `components` and `entities` for UI and interaction systems")
    lines.append("- `server` domain services, validation, and data orchestration")
    lines.append("- `drizzle` and `scripts` for schema and migration complexity")
    lines.append("- `utils/supabase/*`, `supabase/config.toml`, and deployment configuration for hosting scope")
    lines.append("- Informally requested functionality has been formalized and commercialized as payable revision scope")
    lines.append("")
    lines.append("---")
    lines.append("")

    for section in SUMMARY_SECTIONS:
        rows = DATA[section]
        lines.append(f"## {section}")
        lines.append("")
        lines.append("| Task | Description | Roles | Rate/HR | Time | Subtotal |")
        lines.append("|---|---|---|---:|---:|---:|")
        sec_hours = 0
        sec_value = 0
        for task, desc, role, rate, hours in rows:
            resolved_rate = resolve_rate(section, role, rate)
            subtotal = resolved_rate * hours
            sec_hours += hours
            sec_value += subtotal
            lines.append(f"| {task} | {desc} | {role} | {resolved_rate:,.2f} | {hours} | {subtotal:,.2f} |")
        lines.append(f"| **Section Total** |  |  |  | **{sec_hours}** | **{sec_value:,.2f}** |")
        lines.append("")
        lines.append("---")
        lines.append("")

    lines.append("## Final Totals")
    lines.append("")
    lines.append("| Metric | Value |")
    lines.append("|---|---:|")
    for section, (_, value) in section_totals.items():
        lines.append(f"| {section} subtotal | {value:,.2f} |")
    lines.append(f"| **Total Hours** | **{total_hours:,}** |")
    lines.append(f"| **Total Project Value (ex GST)** | **{total_value:,.2f}** |")
    lines.append(f"| Optional GST (10%) | {total_value*0.10:,.2f} |")
    lines.append(f"| Optional Total (inc GST) | {total_value*1.10:,.2f} |")
    lines.append("")
    lines.append("## Assumptions and Exclusions")
    lines.append("")
    lines.append("- This invoice is scoped to the delivery streams represented in the sections above.")
    lines.append("- Any additional out-of-scope functionality should be billed as separate line items or a separate invoice.")
    lines.append("- Third-party pass-through costs (platform subscriptions and transaction fees) are excluded unless listed separately.")

    with open(f"{BASE}/bullyproof-hours-invoice-draft.md", "w", encoding="utf-8") as f:
        f.write("\n".join(lines) + "\n")


if __name__ == "__main__":
    section_totals, total_hours, total_value = compute_totals()
    build_markdown(section_totals, total_hours, total_value)
    build_docx(section_totals, total_hours, total_value, "Bullyproof-Hours-Invoice-Client-Ready.docx")
    build_docx(section_totals, total_hours, total_value, "Bullyproof-Hours-Invoice-Editable.docx")
    build_xlsx(section_totals, total_hours, total_value)
    print(f"Generated assets with {total_hours} hours and A${total_value:,.0f}")
